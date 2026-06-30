from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
MATERIAL = ROOT / "docs" / "课程设计报告素材.md"
OUT = ROOT / "reports" / "电影评分系统课程设计报告.docx"


def set_run_font(run, font_name="宋体", size=Pt(12), bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = size
    run.bold = bold


def set_paragraph_format(paragraph, first_line=True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)


def add_center_paragraph(doc, text, size=Pt(16), bold=False, font="宋体", after=Pt(0)):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = after
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    return paragraph


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", Pt(12))
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", Pt(14 if level == 1 else 12), True)
    return paragraph


def add_code_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", Pt(10))
    return paragraph


def add_numbered_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line=False)
    paragraph.paragraph_format.left_indent = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", Pt(12))
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        paragraph = hdr[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, "宋体", Pt(10.5), True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            paragraph = cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            set_run_font(run, "宋体", Pt(10.5))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_cover(doc):
    add_center_paragraph(doc, "上海电力大学", Pt(18), True, "宋体", Pt(18))
    add_center_paragraph(doc, "数据库应用课程设计", Pt(22), True, "黑体", Pt(44))

    lines = [
        ("题　　目：", "电影评分系统"),
        ("学生姓名：", ""),
        ("学　　号：", ""),
        ("院　　系：", "计算机科学与技术学院"),
        ("专业班级：", ""),
    ]
    for label, value in lines:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.8
        run1 = paragraph.add_run(label)
        set_run_font(run1, "宋体", Pt(14))
        run2 = paragraph.add_run(value if value else "　　　　　　　　　　　　　")
        set_run_font(run2, "宋体", Pt(14))

    doc.add_paragraph()
    add_center_paragraph(doc, "2026 年 6 月 30 日", Pt(14), False, "宋体")
    doc.add_page_break()


def add_data_dictionary(doc):
    add_heading(doc, "数据字典", 2)
    headers = ["表名", "主要字段", "说明"]
    rows = [
        ["users", "user_id, username, password_hash, real_name, role", "保存系统用户及角色信息。"],
        ["movies", "movie_id, title, release_year, genre, rating", "保存电影基本信息和评分统计字段。"],
        ["actors", "actor_id, name, gender, birth_date, nationality", "保存演员基本信息。"],
        ["directors", "director_id, name, gender, birth_date, nationality", "保存导演基本信息。"],
        ["movie_actor", "movie_id, actor_id, role_name", "保存电影与演员的多对多关系。"],
        ["movie_director", "movie_id, director_id", "保存电影与导演的多对多关系。"],
        ["ratings", "rating_id, movie_id, user_id, score", "保存用户对电影的评分。"],
        ["comments", "comment_id, movie_id, user_id, content", "保存用户对电影的评论。"],
    ]
    add_table(doc, headers, rows)


def add_file_table(doc):
    add_heading(doc, "项目文件说明", 2)
    headers = ["文件", "用途"]
    rows = [
        ["server.js", "后端 HTTP 服务、登录认证、权限校验和业务接口。"],
        ["public/app.js", "前端页面渲染、表单校验、API 调用和交互逻辑。"],
        ["public/styles.css", "系统界面样式。"],
        ["data/db.json", "本地演示数据。"],
        ["sql/01_schema.sql", "MySQL 数据库结构、视图、存储过程和触发器。"],
        ["sql/02_seed.sql", "MySQL 测试数据和存储过程调用。"],
    ]
    add_table(doc, headers, rows)


def add_report_body(doc):
    text = MATERIAL.read_text(encoding="utf-8").splitlines()
    in_cover = False
    for line in text:
        raw = line.rstrip()
        if not raw:
            continue
        if raw.startswith("# "):
            continue
        if raw == "## 封面信息":
            in_cover = True
            continue
        if raw.startswith("## 一、"):
            in_cover = False
        if in_cover:
            continue
        if raw.startswith("## "):
            add_heading(doc, raw[3:], 1)
            continue
        if raw.startswith("### "):
            add_heading(doc, raw[4:], 2)
            if raw == "### 2. 关系模式设计":
                add_data_dictionary(doc)
            continue
        if raw.startswith("`") and raw.endswith("`"):
            add_code_paragraph(doc, raw.strip("`"))
            continue
        if raw[:3].strip(".").isdigit() or (len(raw) > 2 and raw[0].isdigit() and raw[1] == "."):
            add_numbered_paragraph(doc, raw)
            continue
        add_body_paragraph(doc, raw)
        if raw == "### 8. 关键文件":
            add_file_table(doc)

    add_file_table(doc)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_report_body(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

